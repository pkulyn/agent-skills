#!/usr/bin/env python3
"""doc2md: 将 PDF/DOCX/URL 转为 Markdown，用于 my-design 内容输入。

用法：
  from doc2md import convert

  md = convert("/path/to/file.pdf")
  md = convert("/path/to/file.docx")
  md = convert("https://example.com/article")

也支持 CLI：
  python3 doc2md.py input.pdf output.md
  python3 doc2md.py https://example.com output.md
"""
import os, sys, tempfile, subprocess


def convert(source, output_path=None):
    """将 PDF/DOCX/URL 转为 Markdown。

    Args:
        source: 文件路径（.pdf/.docx）或 URL
        output_path: 可选，保存到文件

    Returns:
        Markdown 字符串
    """
    if source.startswith("http://") or source.startswith("https://"):
        md = _from_url(source)
    elif source.lower().endswith(".pdf"):
        md = _from_pdf(source)
    elif source.lower().endswith(".docx"):
        md = _from_docx(source)
    else:
        # Try markitdown as fallback
        md = _from_markitdown(source)

    if output_path and md:
        os.makedirs(os.path.dirname(output_path) or ".", exist_ok=True)
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(md)

    return md or ""


def _from_pdf(path):
    """PDF → Markdown。优先用 markitdown，fallback 用 pdf 工具。"""
    # Try markitdown first (better formatting)
    try:
        result = subprocess.run(
            ["markitdown", path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Fallback: use python pdf extraction
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(path)
        parts = []
        for page in doc:
            text = page.get_text()
            if text.strip():
                parts.append(text.strip())
        doc.close()
        return "\n\n---\n\n".join(parts) if parts else ""
    except ImportError:
        pass

    # Last fallback: pdftotext
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", path, "-"],
            capture_output=True, text=True, timeout=30
        )
        if result.returncode == 0:
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    return ""


def _from_docx(path):
    """DOCX → Markdown。"""
    # Try markitdown first
    try:
        result = subprocess.run(
            ["markitdown", path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Fallback: python-docx
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style = para.style.name.lower() if para.style else ""
            if "heading 1" in style:
                parts.append(f"# {text}")
            elif "heading 2" in style:
                parts.append(f"## {text}")
            elif "heading 3" in style:
                parts.append(f"### {text}")
            elif "list" in style:
                parts.append(f"- {text}")
            else:
                parts.append(text)
        
        # Tables
        for table in doc.tables:
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                parts.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    parts.append("| " + " | ".join("---" for _ in cells) + " |")
        
        return "\n\n".join(parts) if parts else ""
    except ImportError:
        pass

    return ""


def _from_url(url):
    """URL → Markdown。"""
    try:
        import urllib.request
        req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
        resp = urllib.request.urlopen(req, timeout=15)
        html = resp.read().decode('utf-8', errors='replace')
    except Exception as e:
        return f"<!-- URL fetch failed: {e} -->"

    # Try markitdown from HTML
    try:
        with tempfile.NamedTemporaryFile(suffix=".html", mode='w', delete=False, encoding='utf-8') as f:
            f.write(html)
            tmp = f.name
        result = subprocess.run(
            ["markitdown", tmp],
            capture_output=True, text=True, timeout=30
        )
        os.unlink(tmp)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass

    # Fallback: simple HTML→text
    import re
    text = re.sub(r'<script[^>]*>.*?</script>', '', html, flags=re.DOTALL)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.DOTALL)
    text = re.sub(r'<[^>]+>', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text[:50000]  # Truncate very long pages


def _from_markitdown(path):
    """Fallback: use markitdown for any file type."""
    try:
        result = subprocess.run(
            ["markitdown", path],
            capture_output=True, text=True, timeout=60
        )
        if result.returncode == 0:
            return result.stdout
    except (FileNotFoundError, subprocess.TimeoutExpired):
        pass
    return ""


# ── CLI ──

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("用法: doc2md.py <input.pdf|.docx|url> [output.md]")
        sys.exit(1)
    
    source = sys.argv[1]
    output = sys.argv[2] if len(sys.argv) > 2 else None
    md = convert(source, output)
    
    if output:
        print(f"Converted → {output} ({len(md)} chars)")
    else:
        print(md[:2000])
        if len(md) > 2000:
            print(f"\n... ({len(md)} chars total)")
